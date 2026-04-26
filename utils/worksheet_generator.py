"""
Victory Partner — Worksheet Generator v0.3
- generate_worksheet_content() : Claude API → text content
- generate_worksheet_pdf()     : text → PDF bytes (fpdf2 + Sarabun font)
- create_worksheet()           : text → Word .docx (legacy)
"""

import anthropic
import os
import re
import urllib.request
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────
# PROMPT TEMPLATES ตามประเภทข้อ
# ─────────────────────────────────────────
QUESTION_TYPE_INSTRUCTIONS = {
    "กากบาท": """สร้างข้อสอบแบบปรนัย (กากบาท) ทั้งหมด
- แต่ละข้อมี 4 ตัวเลือก (ก, ข, ค, ง)
- ระบุคำตอบที่ถูกไว้ในส่วนเฉลยท้ายสุดเท่านั้น
- ในส่วนโจทย์อย่าบอกคำตอบ""",

    "ตอบคำถาม": """สร้างข้อสอบแบบอัตนัย (ตอบคำถาม/แสดงวิธีทำ) ทั้งหมด
- แต่ละข้อมีช่องว่างสำหรับเขียนตอบ (เว้นบรรทัด 3-4 บรรทัด)
- ในส่วนเฉลยท้ายสุดให้เฉลยพร้อมวิธีทำ""",

    "ผสม": """สร้างข้อสอบแบบผสม แบ่งเป็น 2 ส่วน:
ส่วนที่ 1: ปรนัย (กากบาท) ครึ่งหนึ่งของจำนวนข้อ
  - แต่ละข้อมี 4 ตัวเลือก (ก, ข, ค, ง)
ส่วนที่ 2: อัตนัย (แสดงวิธีทำ) ครึ่งหนึ่งของจำนวนข้อ
  - มีช่องว่างสำหรับเขียนตอบ
- เฉลยทั้งสองส่วนอยู่ท้ายสุด""",
}


# ─────────────────────────────────────────
# WORD HELPERS
# ─────────────────────────────────────────
def _add_page_number(paragraph):
    """แทรก field เลขหน้าใน Word"""
    for tag, text in [("begin", None), ("instrText", "PAGE"), ("end", None)]:
        el = OxmlElement("w:fldChar" if tag in ("begin", "end") else "w:instrText")
        if tag in ("begin", "end"):
            el.set(qn("w:fldCharType"), tag)
        else:
            el.text = text
        paragraph._p.append(el)


def _add_watermark(doc: Document, watermark_path: str):
    """แทรกลายน้ำจางๆ ผ่าน header"""
    if not os.path.exists(watermark_path):
        return
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(watermark_path, width=Inches(3.5))


def _add_divider(doc: Document):
    """เส้นแบ่ง"""
    p = doc.add_paragraph("─" * 60)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ─────────────────────────────────────────
# GENERATE CONTENT (Claude API)
# ─────────────────────────────────────────
def generate_worksheet_content(
    subject: str,
    topic: str,
    level_class: str,
    difficulty: str = "ง่าย",
    standard: str = "",
    lecture_notes: str = "",
    num_questions: int = 10,
    time_minutes: int = 30,
    question_type: str = "ผสม",
) -> str:
    """
    ใช้ Claude สร้างเนื้อหา worksheet
    - เฉลยจะอยู่ท้ายสุดเสมอ แยกชัดเจน
    """
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    std_section = f"\nตัวชี้วัด สพฐ: {standard}" if standard else ""
    notes_section = f"\nสิ่งที่ lecture ครอบคลุม: {lecture_notes}" if lecture_notes else ""
    type_instruction = QUESTION_TYPE_INSTRUCTIONS.get(question_type, QUESTION_TYPE_INSTRUCTIONS["ผสม"])

    prompt = f"""คุณเป็น Victory Partner ผู้เชี่ยวชาญด้านการศึกษาไทย

สร้างแบบฝึกหัดสำหรับ:
- วิชา: {subject}
- หัวข้อ: {topic}
- ระดับชั้น: {level_class}
- ระดับความยาก: {difficulty}{std_section}{notes_section}
- จำนวนข้อ: {num_questions} ข้อ
- เวลาทำ: {time_minutes} นาที

อ้างอิงหลักสูตรแกนกลาง สพฐ. พ.ศ. 2551 เป็นหลัก

ประเภทข้อ:
{type_instruction}

โครงสร้าง worksheet ที่ต้องการ:
1. บรรทัดแรก: "เวลา {time_minutes} นาที   คะแนนเต็ม {num_questions} คะแนน"
2. คำสั่ง (1-2 ประโยค อธิบายการตอบ)
3. โจทย์ทั้งหมด (ห้ามแสดงเฉลยในส่วนนี้เด็ดขาด)
4. บรรทัดว่าง 2 บรรทัด
5. บรรทัด: "═══════════════ เฉลย ═══════════════"
6. เฉลยทุกข้อ (พร้อมวิธีทำในข้ออัตนัย)

เขียนเป็นภาษาไทยทั้งหมด ยกเว้นศัพท์เฉพาะวิชาที่จำเป็น"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )

    return message.content[0].text


# ─────────────────────────────────────────
# CREATE WORD DOCUMENT
# ─────────────────────────────────────────
def create_worksheet(
    student_name: str,
    level: str,
    subject: str,
    topic: str,
    difficulty: str,
    standard: str,
    question_type: str,
    num_questions: int,
    time_minutes: int,
    questions_content: str,
    output_path: str,
    logo_top_left: str = None,
    logo_top_right: str = None,
    logo_bottom_left: str = None,
    logo_bottom_right: str = None,
    watermark: str = None,
) -> str:
    """สร้าง Word worksheet ครบรูปแบบ"""

    doc = Document()

    # Margin
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ═══ HEADER ═══
    hdr = doc.sections[0].header
    tbl = hdr.add_table(1, 3, width=Inches(6.5))
    tbl.autofit = False

    # บนซ้าย: logo
    if logo_top_left and os.path.exists(logo_top_left):
        tbl.cell(0, 0).paragraphs[0].add_run().add_picture(logo_top_left, height=Inches(0.45))

    # บนกลาง: ชื่อโรงเรียน
    c = tbl.cell(0, 1).paragraphs[0]
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run("Victory Academy")
    r.bold = True
    r.font.size = Pt(10)

    # บนขวา: logo + เลขหน้า
    c_right = tbl.cell(0, 2).paragraphs[0]
    c_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c_right.add_run("หน้า ")
    _add_page_number(c_right)
    if logo_top_right and os.path.exists(logo_top_right):
        p2 = tbl.cell(0, 2).add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run().add_picture(logo_top_right, height=Inches(0.35))

    # ═══ FOOTER ═══
    ftr = doc.sections[0].footer
    ftbl = ftr.add_table(1, 2, width=Inches(6.5))
    if logo_bottom_left and os.path.exists(logo_bottom_left):
        ftbl.cell(0, 0).paragraphs[0].add_run().add_picture(logo_bottom_left, height=Inches(0.35))
    if logo_bottom_right and os.path.exists(logo_bottom_right):
        p_fr = ftbl.cell(0, 1).paragraphs[0]
        p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fr.add_run().add_picture(logo_bottom_right, height=Inches(0.35))

    # ═══ WATERMARK ═══
    if watermark and os.path.exists(watermark):
        _add_watermark(doc, watermark)

    # ═══ TITLE ═══
    title = doc.add_heading(f"แบบฝึกหัดวิชา{subject}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info line
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"หัวข้อ: {topic}   |   ระดับ: {difficulty}   |   ชั้น: {level}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ตัวชี้วัด
    if standard:
        std_p = doc.add_paragraph()
        std_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        std_run = std_p.add_run(f"ตัวชี้วัด: {standard}")
        std_run.font.size = Pt(9)
        std_run.italic = True
        std_run.font.color.rgb = RGBColor(0x55, 0x55, 0x99)

    # ประเภทข้อ
    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_p.add_run(f"ประเภทข้อ: {question_type}   |   {num_questions} ข้อ   |   {time_minutes} นาที").font.size = Pt(9)

    doc.add_paragraph()

    # ชื่อนักเรียน / วันที่
    name_line = doc.add_paragraph()
    r1 = name_line.add_run("ชื่อ-นามสกุล: ")
    r1.bold = True
    name_line.add_run(f"{student_name}" + " " * 25)
    r2 = name_line.add_run("วันที่: ")
    r2.bold = True
    name_line.add_run("_______________")

    r3 = name_line.add_run("   คะแนนที่ได้: ")
    r3.bold = True
    name_line.add_run(f"_______ / {num_questions}")

    _add_divider(doc)
    doc.add_paragraph()

    # ═══ CONTENT (จาก Claude) ═══
    for line in questions_content.split("\n"):
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(12)

    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────
# PDF GENERATION (fpdf2 + Sarabun font)
# ─────────────────────────────────────────

# Map of characters that Sarabun doesn't support → ASCII fallback
_CHAR_MAP = {
    "✓": "v", "✔": "v",
    "✗": "x", "✘": "x",
    "═": "=", "─": "-", "━": "-", "│": "|", "║": "|",
    "∴": "->", "∵": "<=",
    "→": "->", "←": "<-", "↑": "^", "↓": "v",
    "≤": "<=", "≥": ">=", "≠": "!=",
    "×": "x", "÷": "/",
    "²": "^2", "³": "^3",
    "½": "1/2", "¼": "1/4", "¾": "3/4",
    # strip emojis and box-drawing that won't map cleanly
    "💡": "", "📌": "", "📝": "", "⭐": "*", "🔹": "-", "🔸": "-",
    "┌": "+", "┐": "+", "└": "+", "┘": "+", "├": "+", "┤": "+",
    "┬": "+", "┴": "+", "┼": "+",
}


def _sanitize_for_pdf(text: str) -> str:
    """แทนที่ตัวอักษรที่ Sarabun ไม่รองรับด้วย ASCII fallback"""
    result = []
    for ch in text:
        if ch in _CHAR_MAP:
            result.append(_CHAR_MAP[ch])
        else:
            cp = ord(ch)
            # Keep: ASCII, Thai (U+0E00–U+0E7F), basic Latin extended (U+00A0–U+024F)
            if cp < 128 or 0x0E00 <= cp <= 0x0E7F or 0x00A0 <= cp <= 0x024F:
                result.append(ch)
            # Drop anything else (emojis, box-drawing not mapped above, etc.)
    return "".join(result)


_FONT_DIR     = os.path.join(os.path.dirname(__file__), "..", "assets", "fonts")
_FONT_REGULAR = os.path.join(_FONT_DIR, "Sarabun-Regular.ttf")
_FONT_BOLD    = os.path.join(_FONT_DIR, "Sarabun-Bold.ttf")
_FONT_URLS = {
    _FONT_REGULAR: "https://raw.githubusercontent.com/google/fonts/main/ofl/sarabun/Sarabun-Regular.ttf",
    _FONT_BOLD:    "https://raw.githubusercontent.com/google/fonts/main/ofl/sarabun/Sarabun-Bold.ttf",
}


def _ensure_fonts():
    """Download Sarabun font ถ้ายังไม่มีในเครื่อง (ครั้งแรกเท่านั้น)"""
    os.makedirs(_FONT_DIR, exist_ok=True)
    for path, url in _FONT_URLS.items():
        if not os.path.exists(path):
            urllib.request.urlretrieve(url, path)


def generate_worksheet_pdf(content: str, meta: dict) -> bytes:
    """
    แปลง content text (จาก Claude) เป็น PDF bytes
    meta: {subject, topic, level, standard, num_questions, time_minutes, question_type}
    """
    from fpdf import FPDF

    _ensure_fonts()

    pdf = FPDF(format="A4")
    pdf.set_margins(left=20, top=15, right=20)
    pdf.add_page()
    pdf.add_font("Sarabun", style="",  fname=_FONT_REGULAR)
    pdf.add_font("Sarabun", style="B", fname=_FONT_BOLD)

    W = pdf.epw  # effective page width

    # ── Header ──
    pdf.set_font("Sarabun", "B", 17)
    pdf.cell(W, 11, "Victory Academy", new_x="LMARGIN", new_y="NEXT", align="C")

    pdf.set_font("Sarabun", "B", 13)
    pdf.cell(W, 8, _sanitize_for_pdf(f"{meta.get('subject','')}  ชั้น {meta.get('level','')}"),
             new_x="LMARGIN", new_y="NEXT", align="C")

    pdf.set_font("Sarabun", "", 12)
    pdf.cell(W, 7, _sanitize_for_pdf(f"หัวข้อ: {meta.get('topic','')}"),
             new_x="LMARGIN", new_y="NEXT", align="C")

    if meta.get("standard"):
        pdf.set_font("Sarabun", "", 10)
        pdf.multi_cell(W, 6, _sanitize_for_pdf(f"ตัวชี้วัด: {meta.get('standard','')}"), align="C")

    pdf.set_font("Sarabun", "", 11)
    pdf.cell(W, 6,
             f"เวลา {meta.get('time_minutes', 30)} นาที  |  "
             f"{meta.get('num_questions', 10)} ข้อ  |  "
             f"{meta.get('question_type', '')}",
             new_x="LMARGIN", new_y="NEXT", align="C")

    # divider
    pdf.ln(3)
    pdf.set_draw_color(150, 150, 150)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(6)

    # ── Name / Date / Score ──
    pdf.set_font("Sarabun", "", 12)
    pdf.cell(W * 0.5, 8, "ชื่อ-นามสกุล .............................................",
             new_x="RIGHT", new_y="TOP")
    pdf.cell(W * 0.3, 8, "วันที่ ..................................",
             new_x="RIGHT", new_y="TOP")
    pdf.cell(W * 0.2, 8, f"คะแนน ..../{meta.get('num_questions',10)}",
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # ── Content ──
    pdf.set_font("Sarabun", "", 12)

    for line in content.split("\n"):
        raw = line.rstrip()

        if not raw:
            pdf.ln(2)
            continue

        # ตรวจก่อน sanitize เพราะ ═ จะถูกแปลงเป็น = แล้ว
        is_answer_header = "เฉลย" in raw and (raw.count("═") >= 3 or raw.count("=") >= 3)

        s = _sanitize_for_pdf(raw)

        # เฉลย section header
        if is_answer_header:
            pdf.ln(5)
            pdf.set_font("Sarabun", "B", 12)
            pdf.set_fill_color(220, 220, 220)
            pdf.cell(W, 8, "  === เฉลย ===", fill=True,
                     new_x="LMARGIN", new_y="NEXT")
            pdf.set_fill_color(255, 255, 255)
            pdf.set_font("Sarabun", "", 11)
            pdf.ln(2)
            continue

        # Question number (1. 2. ...)
        if re.match(r"^\d+[\.\)]", s):
            pdf.ln(2)
            pdf.set_font("Sarabun", "B", 12)
            pdf.multi_cell(W, 7, s)
            pdf.set_font("Sarabun", "", 12)
            continue

        # MCQ options ก. ข. ค. ง.
        if re.match(r"^[กขคง][\.\)]", s):
            pdf.set_x(pdf.l_margin + 8)
            pdf.multi_cell(W - 8, 6.5, s)
            continue

        # ส่วนหัวข้อ / คำสั่ง (ตัวหนา)
        if s.startswith("ส่วนที่") or s.startswith("คำสั่ง") or s.startswith("คำชี้แจง"):
            pdf.set_font("Sarabun", "B", 12)
            pdf.multi_cell(W, 7, s)
            pdf.set_font("Sarabun", "", 12)
            continue

        # regular text
        pdf.multi_cell(W, 7, s)

    return bytes(pdf.output())
